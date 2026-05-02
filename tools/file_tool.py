"""文件处理工具 - 文件验证、文本提取、预处理"""

import os
from typing import Optional, Tuple

from config.config import get_config
from storage.log_manager import get_logger
from utils.file_utils import (
    get_file_extension,
    is_supported_file,
    check_file_size,
)
from utils.text_utils import truncate_text, clean_text



logger = get_logger("FileTool")


def extract_text_from_pdf(filepath: str) -> Tuple[bool, str]:
    """从PDF文件中提取文本

    Args:
        filepath: PDF文件路径

    Returns:
        (是否成功, 提取的文本或错误信息)
    """
    try:
        import PyPDF2

        text_parts = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)

            logger.info(f"PDF文件页数: {num_pages}")

            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"提取第{i+1}页文本失败: {e}")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            return False, "PDF文件中未找到可提取的文本内容"

        logger.info(f"PDF文本提取成功: {len(full_text)}字符")
        return True, full_text

    except ImportError:
        return False, "缺少PyPDF2依赖库，请安装: pip install PyPDF2"
    except Exception as e:
        logger.error(f"PDF文本提取失败: {e}", exc_info=True)
        return False, f"PDF文件解析失败: {str(e)}"


def extract_text_from_docx(filepath: str) -> Tuple[bool, str]:
    """从Word文档中提取文本

    Args:
        filepath: Word文档路径

    Returns:
        (是否成功, 提取的文本或错误信息)
    """
    try:
        from docx import Document

        doc = Document(filepath)
        text_parts = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            return False, "Word文档中未找到可提取的文本内容"

        logger.info(f"Word文档文本提取成功: {len(full_text)}字符")
        return True, full_text

    except ImportError:
        return False, "缺少python-docx依赖库，请安装: pip install python-docx"
    except Exception as e:
        logger.error(f"Word文档文本提取失败: {e}", exc_info=True)
        return False, f"Word文档解析失败: {str(e)}"


def extract_text_from_txt(filepath: str) -> Tuple[bool, str]:
    """从文本文件中读取内容

    Args:
        filepath: 文本文件路径

    Returns:
        (是否成功, 提取的文本或错误信息)
    """
    try:
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        for encoding in encodings:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    text = f.read()
                    logger.info(f"文本文件读取成功 (编码: {encoding}): {len(text)}字符")
                    return True, text
            except UnicodeDecodeError:
                continue

        return False, "无法识别文件编码，请使用UTF-8编码的文件"

    except Exception as e:
        logger.error(f"文本文件读取失败: {e}", exc_info=True)
        return False, f"文件读取失败: {str(e)}"


def process_file(filepath: str, filename: Optional[str] = None) -> Tuple[bool, str, str]:
    """处理文件，提取文本

    Args:
        filepath: 文件路径
        filename: 文件名（用于确定文件类型）

    Returns:
        (是否成功, 提取的文本或错误信息, 原始文件名)
    """
    if not filename:
        filename = os.path.basename(filepath)

    config = get_config()

    logger.info(f"开始处理文件: {filename}")

    # 检查文件是否存在
    if not os.path.exists(filepath):
        return False, f"文件不存在: {filepath}", filename

    # 检查文件大小
    file_size = os.path.getsize(filepath)
    size_ok, size_msg = check_file_size(file_size, config.max_file_size)
    if not size_ok:
        return False, size_msg, filename

    # 检查文件格式
    if not is_supported_file(filename):
        return False, f"不支持的文件格式: {filename}", filename

    # 根据扩展名选择提取方法
    ext = get_file_extension(filename).lower()

    success = False
    text = ""
    error_msg = ""

    if ext == "pdf":
        success, text = extract_text_from_pdf(filepath)
    elif ext == "docx":
        success, text = extract_text_from_docx(filepath)
    elif ext == "txt":
        success, text = extract_text_from_txt(filepath)
    else:
        error_msg = f"不支持的文件格式: {ext}"
        logger.error(error_msg)
        return False, error_msg, filename

    if not success:
        return False, text, filename

    # 清洗和截断文本
    text = clean_text(text)

    if len(text) > config.max_text_length:
        logger.warning(f"文本过长（{len(text)}字符），将截断为{config.max_text_length}字符")
        text = truncate_text(text, config.max_text_length)

    logger.info(f"文件处理完成: {filename}, 文本长度: {len(text)}字符")
    return True, text, filename


def validate_file(
    filename: str,
    file_size: int,
) -> Tuple[bool, str]:
    """验证文件是否支持

    Args:
        filename: 文件名
        file_size: 文件大小（字节）

    Returns:
        (是否有效, 错误信息)
    """
    config = get_config()

    # 检查格式
    if not is_supported_file(filename):
        return False, f"不支持的文件格式: {filename}\n支持的格式: PDF, DOCX, TXT"

    # 检查大小
    size_ok, size_msg = check_file_size(file_size, config.max_file_size)
    if not size_ok:
        return False, size_msg

    return True, ""
